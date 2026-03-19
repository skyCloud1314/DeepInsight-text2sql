"""
Intel DeepInsight 导出和分享管理器
支持PDF报告生成、会话分享和数据导出功能
包含完整对话内容、图表、AI思考过程等的综合导出
"""
import pandas as pd
import json
import os
from datetime import datetime
from typing import Dict, List, Optional, Any
import base64
from io import BytesIO
import uuid

# 尝试导入PDF生成库
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("⚠️ 提示: 未安装 reportlab 库，PDF 导出功能将不可用。运行: pip install reportlab")

# 尝试导入图表处理库
try:
    import plotly.graph_objects as go
    import plotly.io as pio
    from PIL import Image as PILImage
    CHART_EXPORT_AVAILABLE = True
except ImportError:
    CHART_EXPORT_AVAILABLE = False
    print("⚠️ 提示: 图表导出功能需要 plotly 和 PIL 库")

# 尝试导入matplotlib作为备用图表库
try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # 使用非交互式后端
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    print("⚠️ 提示: matplotlib 不可用，图表导出功能受限")

# 尝试导入DOCX处理库
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ 提示: 未安装 python-docx 库，DOCX 导出功能将不可用。运行: pip install python-docx")

class ExportManager:
    """导出和分享管理器"""
    
    def __init__(self):
        self.exports_dir = "data/exports"
        self.shares_dir = "data/shares"
        self._ensure_directories()
        self._setup_chinese_fonts()
    
    def _ensure_directories(self):
        """确保导出目录存在"""
        os.makedirs(self.exports_dir, exist_ok=True)
        os.makedirs(self.shares_dir, exist_ok=True)
    
    def _setup_chinese_fonts(self):
        """设置中文字体支持"""
        if not REPORTLAB_AVAILABLE:
            return
        
        try:
            # 优先使用本地字体文件夹
            local_font_dir = "fonts"
            local_font_paths = []
            
            if os.path.exists(local_font_dir):
                # 扫描本地字体文件夹
                for font_file in os.listdir(local_font_dir):
                    if font_file.lower().endswith(('.ttf', '.otf', '.ttc')):
                        local_font_paths.append(os.path.join(local_font_dir, font_file))
            
            # 如果本地字体文件夹有字体，优先使用
            if local_font_paths:
                for font_path in local_font_paths:
                    try:
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                        self.chinese_font_available = True
                        print(f"✅ 成功注册本地中文字体: {font_path}")
                        return
                    except Exception as e:
                        print(f"⚠️ 本地字体注册失败: {font_path}, 错误: {e}")
                        continue
            
            # 如果本地字体不可用，回退到系统字体
            print("📁 本地字体不可用，尝试系统字体...")
            import platform
            system = platform.system()
            
            if system == "Windows":
                # Windows系统字体路径
                font_paths = [
                    "C:/Windows/Fonts/simsun.ttc",  # 宋体
                    "C:/Windows/Fonts/simhei.ttf",  # 黑体
                    "C:/Windows/Fonts/msyh.ttc",    # 微软雅黑
                ]
            elif system == "Darwin":  # macOS
                font_paths = [
                    "/System/Library/Fonts/PingFang.ttc",
                    "/System/Library/Fonts/STHeiti Light.ttc",
                ]
            else:  # Linux
                font_paths = [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
                    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",  # 常见的Linux中文字体
                ]
            
            # 尝试注册第一个可用的系统字体
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                        self.chinese_font_available = True
                        print(f"✅ 成功注册系统中文字体: {font_path}")
                        return
                    except Exception as e:
                        continue
            
            # 如果都没有找到，创建默认字体提示
            self.chinese_font_available = False
            print("⚠️ 未找到中文字体，PDF中的中文可能显示为方块")
            print("💡 建议：将中文字体文件（.ttf/.otf）放入 fonts/ 文件夹中")
            
        except Exception as e:
            self.chinese_font_available = False
            print(f"❌ 字体初始化失败: {e}")
            print(f"字体设置失败: {e}")
    
    def _convert_chart_to_image(self, chart_data: Dict) -> Optional[str]:
        """将图表数据转换为图片文件，支持多种导出引擎 - 改进的错误处理和超时机制"""
        if not (CHART_EXPORT_AVAILABLE or MATPLOTLIB_AVAILABLE):
            print("⚠️ 图表导出功能不可用：缺少matplotlib和plotly库")
            return None
        
        try:
            chart_type = chart_data.get("type", "bar")
            data = chart_data.get("data", {})
            title = chart_data.get("title", "数据图表")
            
            if not data:
                print("⚠️ 图表数据为空，跳过图表生成")
                return None
            
            # 生成文件路径
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            img_filename = f"chart_{timestamp}.png"
            img_path = os.path.join(self.exports_dir, img_filename)
            
            print(f"🎨 开始生成图表: {title} ({chart_type})")
            
            # 首先尝试使用matplotlib（更可靠，无外部依赖）
            if MATPLOTLIB_AVAILABLE:
                try:
                    print("   尝试使用matplotlib引擎...")
                    success = self._create_chart_with_matplotlib(chart_type, data, title, img_path)
                    if success and os.path.exists(img_path) and os.path.getsize(img_path) > 0:
                        print(f"   ✅ matplotlib生成成功: {os.path.basename(img_path)}")
                        return img_path
                    else:
                        print("   ❌ matplotlib生成失败")
                except Exception as e:
                    print(f"   ❌ matplotlib异常: {e}")
            
            # 如果matplotlib失败，尝试plotly（功能更丰富但可能有依赖问题）
            if CHART_EXPORT_AVAILABLE:
                try:
                    print("   尝试使用plotly引擎...")
                    success = self._create_chart_with_plotly(chart_type, data, title, img_path)
                    if success and os.path.exists(img_path) and os.path.getsize(img_path) > 0:
                        print(f"   ✅ plotly生成成功: {os.path.basename(img_path)}")
                        return img_path
                    else:
                        print("   ❌ plotly生成失败")
                except Exception as e:
                    print(f"   ❌ plotly异常: {e}")
            
            print(f"   ❌ 所有图表引擎都失败，无法生成图表: {title}")
            return None
            
        except Exception as e:
            print(f"❌ 图表转换失败: {e}")
            return None
    
    def _create_chart_with_matplotlib(self, chart_type: str, data: Dict, title: str, img_path: str) -> bool:
        """使用matplotlib创建图表"""
        try:
            plt.figure(figsize=(8, 6))
            plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
            plt.rcParams['axes.unicode_minus'] = False
            
            if chart_type == "bar":
                x_data = data.get("x", [])
                y_data = data.get("y", [])
                plt.bar(x_data, y_data, color='#0068B5', alpha=0.8)
                plt.xlabel("类别")
                plt.ylabel("数值")
                
            elif chart_type == "line":
                x_data = data.get("x", [])
                y_data = data.get("y", [])
                plt.plot(x_data, y_data, marker='o', linewidth=2, color='#0068B5')
                plt.xlabel("类别")
                plt.ylabel("数值")
                
            elif chart_type == "pie":
                labels = data.get("labels", [])
                values = data.get("values", [])
                colors = ['#0068B5', '#00a8ff', '#74b9ff', '#a29bfe', '#fd79a8']
                plt.pie(values, labels=labels, autopct='%1.1f%%', colors=colors[:len(values)])
                
            elif chart_type == "scatter":
                x_data = data.get("x", [])
                y_data = data.get("y", [])
                plt.scatter(x_data, y_data, color='#0068B5', alpha=0.7, s=60)
                plt.xlabel("X轴")
                plt.ylabel("Y轴")
            
            plt.title(title, fontsize=14, fontweight='bold', pad=20)
            plt.tight_layout()
            plt.savefig(img_path, dpi=150, bbox_inches='tight', facecolor='white')
            plt.close()
            
            return True
            
        except Exception as e:
            print(f"Matplotlib图表创建失败: {e}")
            plt.close()
            return False
    
    def _create_chart_with_plotly(self, chart_type: str, data: Dict, title: str, img_path: str) -> bool:
        """使用plotly创建图表（备用方案）- 改进的超时处理"""
        try:
            # 创建图表
            if chart_type == "bar":
                fig = go.Figure(data=[
                    go.Bar(x=data.get("x", []), y=data.get("y", []), name=data.get("name", ""))
                ])
            elif chart_type == "line":
                fig = go.Figure(data=[
                    go.Scatter(x=data.get("x", []), y=data.get("y", []), mode='lines+markers', name=data.get("name", ""))
                ])
            elif chart_type == "pie":
                fig = go.Figure(data=[
                    go.Pie(labels=data.get("labels", []), values=data.get("values", []))
                ])
            elif chart_type == "scatter":
                fig = go.Figure(data=[
                    go.Scatter(x=data.get("x", []), y=data.get("y", []), mode='markers', name=data.get("name", ""))
                ])
            else:
                return False
            
            # 设置图表布局
            fig.update_layout(
                title=title,
                width=600,
                height=400,
                font=dict(size=10),
                margin=dict(l=50, r=50, t=50, b=50)
            )
            
            # 改进的导出引擎尝试机制 - 更短的超时时间和更好的错误处理
            engines_to_try = ['kaleido', 'orca']
            
            for engine in engines_to_try:
                try:
                    import threading
                    import time
                    
                    result = [False]
                    exception = [None]
                    start_time = time.time()
                    
                    def export_with_timeout():
                        try:
                            # 设置引擎特定的超时
                            import plotly.io as pio_local
                            pio_local.write_image(fig, img_path, format='png', engine=engine, timeout=3)
                            result[0] = True
                        except Exception as e:
                            exception[0] = e
                    
                    # 创建线程执行导出，使用更短的超时时间
                    thread = threading.Thread(target=export_with_timeout)
                    thread.daemon = True
                    thread.start()
                    thread.join(timeout=3)  # 减少到3秒超时
                    
                    elapsed_time = time.time() - start_time
                    
                    if thread.is_alive():
                        print(f"Plotly引擎 {engine} 超时 ({elapsed_time:.1f}s)")
                        # 强制结束线程（通过设置daemon=True已经处理）
                        continue
                    
                    if exception[0]:
                        print(f"Plotly引擎 {engine} 失败: {exception[0]}")
                        continue
                    
                    if result[0] and os.path.exists(img_path):
                        print(f"Plotly引擎 {engine} 成功 ({elapsed_time:.1f}s)")
                        return True
                        
                except Exception as e:
                    print(f"尝试Plotly引擎 {engine} 时发生错误: {e}")
                    continue
            
            print("所有Plotly引擎都失败，将使用matplotlib备用方案")
            return False
            
        except Exception as e:
            print(f"Plotly图表创建失败: {e}")
            return False
    
    def _add_chart_to_story(self, story: List, chart_data: Dict, styles: Dict):
        """将图表添加到PDF故事中"""
        try:
            img_path = self._convert_chart_to_image(chart_data)
            if img_path and os.path.exists(img_path):
                # 添加图表标题
                chart_title = chart_data.get("title", "数据图表")
                story.append(Paragraph(f"📊 {chart_title}", styles["heading4_style"]))
                
                # 读取图片数据并创建BytesIO对象
                try:
                    with open(img_path, 'rb') as f:
                        img_data = f.read()
                    
                    from io import BytesIO
                    img_buffer = BytesIO(img_data)
                    img = Image(img_buffer, width=5*inch, height=3.3*inch)
                    story.append(img)
                    story.append(Spacer(1, 10))
                    
                    # 记录需要清理的文件
                    if not hasattr(self, '_temp_chart_files'):
                        self._temp_chart_files = []
                    self._temp_chart_files.append(img_path)
                    
                except Exception as e:
                    print(f"读取图片文件失败: {e}")
                    story.append(Paragraph("📊 [图表数据 - 读取失败]", styles["normal_style"]))
                    story.append(Spacer(1, 10))
            else:
                # 如果图表转换失败，添加占位符
                story.append(Paragraph("📊 [图表数据 - 转换失败]", styles["normal_style"]))
                story.append(Spacer(1, 10))
        except Exception as e:
            print(f"添加图表到PDF失败: {e}")
            story.append(Paragraph("📊 [图表数据 - 处理失败]", styles["normal_style"]))
            story.append(Spacer(1, 10))
    def export_session_to_pdf(self, session_data: Dict, session_title: str = "分析报告") -> Optional[str]:
        """导出完整会话为PDF报告，包含所有对话内容、图表、AI思考过程等"""
        if not REPORTLAB_AVAILABLE:
            return None
        
        try:
            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"DeepInsight_Complete_Report_{timestamp}.pdf"
            filepath = os.path.join(self.exports_dir, filename)
            
            # 创建PDF文档
            doc = SimpleDocTemplate(filepath, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
            styles = getSampleStyleSheet()
            story = []
            
            # 设置样式字典
            style_dict = self._setup_pdf_styles(styles)
            
            # 标题页
            story.append(Paragraph("Intel® DeepInsight 智能分析报告", style_dict["title_style"]))
            story.append(Spacer(1, 20))
            story.append(Paragraph(f"会话标题: {session_title}", style_dict["heading_style"]))
            story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}", style_dict["normal_style"]))
            story.append(Paragraph(f"系统版本: Intel® DeepInsight Pro v2.0", style_dict["normal_style"]))
            story.append(Spacer(1, 30))
            
            # 目录概览
            story.append(Paragraph("📋 报告内容概览", style_dict["heading_style"]))
            messages = session_data.get("messages", [])
            question_count = len([msg for msg in messages if msg["role"] == "user"])
            story.append(Paragraph(f"• 用户问题数量: {question_count}", style_dict["normal_style"]))
            story.append(Paragraph(f"• AI回答数量: {len([msg for msg in messages if msg['role'] == 'assistant'])}", style_dict["normal_style"]))
            story.append(Paragraph(f"• 包含SQL查询: {len([msg for msg in messages if msg.get('sql')])}", style_dict["normal_style"]))
            story.append(Paragraph(f"• 包含AI思考过程: {len([msg for msg in messages if msg.get('thought')])}", style_dict["normal_style"]))
            story.append(Spacer(1, 30))
            
            # 处理每个对话
            question_num = 1
            for i, msg in enumerate(messages):
                if msg["role"] == "user":
                    # 用户问题
                    story.append(Paragraph(f"[*] 问题 {question_num}: {msg['content']}", style_dict["heading3_style"]))
                    story.append(Spacer(1, 10))
                    question_num += 1
                
                elif msg["role"] == "assistant":
                    # AI思考过程 (如果有) - 使用用户配置的模型名称
                    if msg.get("thought"):
                        # 获取用户配置的模型名称
                        model_name = "AI模型"  # 默认值
                        try:
                            import streamlit as st
                            if hasattr(st, 'session_state') and hasattr(st.session_state, 'config'):
                                model_name = st.session_state.config.get("model_name", "AI模型")
                        except:
                            pass
                        
                        story.append(Paragraph(f"🧠 AI思考过程 ({model_name})", style_dict["heading4_style"]))
                        thought_content = msg["thought"][:1000] + "..." if len(msg["thought"]) > 1000 else msg["thought"]
                        if hasattr(self, 'chinese_font_available') and self.chinese_font_available:
                            story.append(Paragraph(thought_content, style_dict["thought_style"]))
                        else:
                            story.append(Paragraph("[AI思考过程 - 需要中文字体支持]", style_dict["thought_style"]))
                        # 思考过程后的间距已在样式中定义，无需额外添加
                    
                    # 表选择过程信息 (如果有)
                    self._add_table_selection_info(story, msg.get("table_selection_info"), style_dict)
                    
                    # SQL查询 (如果有)
                    if msg.get("sql"):
                        story.append(Paragraph("💻 生成的SQL查询", style_dict["heading4_style"]))
                        story.append(Paragraph(msg["sql"], style_dict["sql_style"]))
                        story.append(Spacer(1, 10))
                    
                    # 查询结果和数据
                    if "data" in msg and msg["data"]:
                        story.append(Paragraph("📊 查询结果", style_dict["heading4_style"]))
                        
                        # 数据表格
                        self._add_data_table(story, msg["data"], style_dict)
                        
                        # 如果有图表数据，添加图表
                        if msg.get("charts"):
                            for chart_data in msg["charts"]:
                                self._add_chart_to_story(story, chart_data, style_dict)
                    
                    # 商业洞察和分析 - 添加适当间距
                    if msg.get("content"):
                        story.append(Spacer(1, 35))  # 与上一个内容块的间距
                        story.append(Paragraph("💡 商业洞察与分析", style_dict["heading4_style"]))
                        content = msg["content"]
                        if hasattr(self, 'chinese_font_available') and self.chinese_font_available:
                            story.append(Paragraph(content, style_dict["normal_style"]))
                        else:
                            story.append(Paragraph("[商业洞察内容 - 需要中文字体支持]", style_dict["normal_style"]))
                        story.append(Spacer(1, 20))  # 与理解方式部分的间距
                    
                    # 其他可能的理解方式和选中的可能性 - 使用修复后的方法
                    self._add_alternatives_info(story, msg, style_dict)
                    
                    story.append(Spacer(1, 25))  # 对话间的间距
                    # 添加分隔线
                    story.append(Paragraph("─" * 80, style_dict["separator_style"]))
                    story.append(Spacer(1, 20))
            
            # 报告总结
            self._add_report_summary(story, question_count, style_dict)
            
            # 生成PDF
            doc.build(story)
            
            # 清理临时图片文件
            if hasattr(self, '_temp_chart_files'):
                for temp_file in self._temp_chart_files:
                    try:
                        if os.path.exists(temp_file):
                            os.remove(temp_file)
                    except:
                        pass
                delattr(self, '_temp_chart_files')
            
            return filepath
            
        except Exception as e:
            print(f"PDF生成失败: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _setup_pdf_styles(self, styles):
        """设置PDF样式 - 修复视觉重叠问题"""
        style_dict = {}
        
        if hasattr(self, 'chinese_font_available') and self.chinese_font_available:
            # 使用中文字体
            style_dict["title_style"] = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                spaceAfter=30,
                textColor=colors.HexColor('#0068B5'),
                fontName='ChineseFont',
                alignment=1  # 居中
            )
            
            style_dict["heading_style"] = ParagraphStyle(
                'ChineseHeading2',
                parent=styles['Heading2'],
                fontName='ChineseFont',
                fontSize=14,
                spaceAfter=12,
                textColor=colors.HexColor('#0068B5')
            )
            
            style_dict["heading3_style"] = ParagraphStyle(
                'ChineseHeading3',
                parent=styles['Heading3'],
                fontName='ChineseFont',
                fontSize=12,
                spaceAfter=8,
                textColor=colors.HexColor('#2c3e50')
            )
            
            style_dict["heading4_style"] = ParagraphStyle(
                'ChineseHeading4',
                parent=styles['Heading4'],
                fontName='ChineseFont',
                fontSize=10,
                spaceAfter=6,
                textColor=colors.HexColor('#34495e')
            )
            
            style_dict["normal_style"] = ParagraphStyle(
                'ChineseNormal',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=9,
                spaceAfter=6
            )
            
            style_dict["sql_style"] = ParagraphStyle(
                'ChineseSQLStyle',
                parent=styles['Code'],
                fontSize=8,
                leftIndent=20,
                fontName='ChineseFont',
                backColor=colors.HexColor('#f8f9fa'),
                borderColor=colors.HexColor('#dee2e6'),
                borderWidth=1,
                borderPadding=8
            )
            
            # 修复思考过程样式 - 增加间距避免重叠
            style_dict["thought_style"] = ParagraphStyle(
                'ChineseThoughtStyle',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=8,
                leftIndent=15,
                rightIndent=15,
                backColor=colors.HexColor('#f0f7ff'),
                borderColor=colors.HexColor('#0068B5'),
                borderWidth=1,
                borderPadding=12,
                spaceAfter=40,  # 增加到40像素
                spaceBefore=12  # 添加前间距
            )
            
            # 新增：其他理解方式标题样式（灰色主题）
            style_dict["alternatives_heading_style"] = ParagraphStyle(
                'ChineseAlternativesHeadingStyle',
                parent=styles['Heading4'],
                fontName='ChineseFont',
                fontSize=10,
                textColor=colors.HexColor('#6c757d'),
                spaceAfter=8,
                spaceBefore=12
            )
            
            # 新增：其他理解方式内容样式
            style_dict["alternatives_style"] = ParagraphStyle(
                'ChineseAlternativesStyle',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=9,
                leftIndent=10,
                textColor=colors.HexColor('#6c757d'),
                spaceAfter=3
            )
            
            # 新增：选中理解方式标题样式（绿色主题）
            style_dict["selected_heading_style"] = ParagraphStyle(
                'ChineseSelectedHeadingStyle',
                parent=styles['Heading4'],
                fontName='ChineseFont',
                fontSize=10,
                textColor=colors.HexColor('#28a745'),
                spaceAfter=8,
                spaceBefore=12
            )
            
            # 新增：选中理解方式内容样式
            style_dict["selected_style"] = ParagraphStyle(
                'ChineseSelectedStyle',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=9,
                leftIndent=10,
                textColor=colors.HexColor('#28a745'),
                backColor=colors.HexColor('#f8fff8'),
                borderColor=colors.HexColor('#28a745'),
                borderWidth=1,
                borderPadding=8,
                spaceAfter=10
            )
            
            style_dict["table_info_style"] = ParagraphStyle(
                'ChineseTableInfoStyle',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=8,
                leftIndent=10,
                backColor=colors.HexColor('#f8f9fa'),
                borderColor=colors.HexColor('#28a745'),
                borderWidth=1,
                borderPadding=8,
                spaceAfter=6
            )
            
            # 新增：分隔符样式
            style_dict["separator_style"] = ParagraphStyle(
                'ChineseSeparatorStyle',
                parent=styles['Normal'],
                fontName='ChineseFont',
                fontSize=8,
                textColor=colors.HexColor('#dee2e6'),
                alignment=1,  # 居中
                spaceAfter=15,
                spaceBefore=15
            )
        else:
            # 使用默认字体（可能不支持中文）
            style_dict["title_style"] = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                spaceAfter=30,
                textColor=colors.HexColor('#0068B5'),
                alignment=1
            )
            style_dict["heading_style"] = styles['Heading2']
            style_dict["heading3_style"] = styles['Heading3']
            style_dict["heading4_style"] = styles['Heading4']
            style_dict["normal_style"] = styles['Normal']
            style_dict["sql_style"] = ParagraphStyle(
                'SQLStyle',
                parent=styles['Code'],
                fontSize=8,
                leftIndent=20,
                backColor=colors.HexColor('#f8f9fa')
            )
            
            # 修复思考过程样式 - 增加间距避免重叠
            style_dict["thought_style"] = ParagraphStyle(
                'ThoughtStyle',
                parent=styles['Normal'],
                fontSize=8,
                leftIndent=15,
                rightIndent=15,
                backColor=colors.HexColor('#f0f7ff'),
                borderColor=colors.HexColor('#0068B5'),
                borderWidth=1,
                borderPadding=12,
                spaceAfter=40,  # 增加到40像素
                spaceBefore=12  # 添加前间距
            )
            
            # 新增：其他理解方式标题样式（灰色主题）
            style_dict["alternatives_heading_style"] = ParagraphStyle(
                'AlternativesHeadingStyle',
                parent=styles['Heading4'],
                fontSize=10,
                textColor=colors.HexColor('#6c757d'),
                spaceAfter=8,
                spaceBefore=12
            )
            
            # 新增：其他理解方式内容样式
            style_dict["alternatives_style"] = ParagraphStyle(
                'AlternativesStyle',
                parent=styles['Normal'],
                fontSize=9,
                leftIndent=10,
                textColor=colors.HexColor('#6c757d'),
                spaceAfter=3
            )
            
            # 新增：选中理解方式标题样式（绿色主题）
            style_dict["selected_heading_style"] = ParagraphStyle(
                'SelectedHeadingStyle',
                parent=styles['Heading4'],
                fontSize=10,
                textColor=colors.HexColor('#28a745'),
                spaceAfter=8,
                spaceBefore=12
            )
            
            # 新增：选中理解方式内容样式
            style_dict["selected_style"] = ParagraphStyle(
                'SelectedStyle',
                parent=styles['Normal'],
                fontSize=9,
                leftIndent=10,
                textColor=colors.HexColor('#28a745'),
                backColor=colors.HexColor('#f8fff8'),
                borderColor=colors.HexColor('#28a745'),
                borderWidth=1,
                borderPadding=8,
                spaceAfter=10
            )
            
            style_dict["table_info_style"] = ParagraphStyle(
                'TableInfoStyle',
                parent=styles['Normal'],
                fontSize=8,
                leftIndent=10,
                backColor=colors.HexColor('#f8f9fa')
            )
            
            # 新增：分隔符样式
            style_dict["separator_style"] = ParagraphStyle(
                'SeparatorStyle',
                parent=styles['Normal'],
                fontSize=8,
                textColor=colors.HexColor('#dee2e6'),
                alignment=1,  # 居中
                spaceAfter=15,
                spaceBefore=15
            )
        
        return style_dict
    
    def _add_table_selection_info(self, story: List, table_info: Dict, style_dict: Dict):
        """添加表选择过程信息到PDF"""
        if not table_info or not any(table_info.values()):
            return
        
        story.append(Paragraph("🗄️ 智能表选择过程", style_dict["heading4_style"]))
        
        # 初步分析
        if table_info.get("initial_analysis"):
            story.append(Paragraph("第1步: 语义相似度初步筛选", style_dict["heading4_style"]))
            if hasattr(self, 'chinese_font_available') and self.chinese_font_available:
                story.append(Paragraph(table_info["initial_analysis"], style_dict["table_info_style"]))
            else:
                story.append(Paragraph("[表选择分析 - 需要中文字体支持]", style_dict["table_info_style"]))
        
        # Agent推理
        if table_info.get("agent_reasoning"):
            story.append(Paragraph("第2步: Agent智能筛选推理", style_dict["heading4_style"]))
            if hasattr(self, 'chinese_font_available') and self.chinese_font_available:
                story.append(Paragraph(f"推理过程: {table_info['agent_reasoning']}", style_dict["table_info_style"]))
            else:
                story.append(Paragraph("[Agent推理过程 - 需要中文字体支持]", style_dict["table_info_style"]))
        
        # 关联分析
        if table_info.get("join_analysis"):
            story.append(Paragraph("第3步: 表关联关系分析", style_dict["heading4_style"]))
            if hasattr(self, 'chinese_font_available') and self.chinese_font_available:
                story.append(Paragraph(table_info["join_analysis"], style_dict["table_info_style"]))
            else:
                story.append(Paragraph("[关联分析 - 需要中文字体支持]", style_dict["table_info_style"]))
        
        # 最终选择结果
        if table_info.get("final_selection"):
            final_selection = table_info["final_selection"]
            selected_tables = final_selection.get("selected_tables", [])
            if selected_tables:
                story.append(Paragraph("最终选择的表:", style_dict["heading4_style"]))
                for table in selected_tables:
                    table_name = table.get("table_name", "未知表")
                    relevance_score = table.get("relevance_score", 0)
                    reasoning = table.get("reasoning", "无推理信息")
                    table_text = f"• {table_name} (相关度: {relevance_score:.2f}) - {reasoning}"
                    if hasattr(self, 'chinese_font_available') and self.chinese_font_available:
                        story.append(Paragraph(table_text, style_dict["normal_style"]))
                    else:
                        story.append(Paragraph(f"• {table_name} (Score: {relevance_score:.2f})", style_dict["normal_style"]))
        
        story.append(Spacer(1, 10))
    
    def _add_data_table(self, story: List, data: List, style_dict: Dict):
        """添加数据表格到PDF"""
        df = pd.DataFrame(data)
        if not df.empty:
            # 限制表格大小以适应PDF
            max_rows = 15
            max_cols = 8
            
            # 截取数据
            display_df = df.iloc[:max_rows, :max_cols]
            
            # 转换数据为字符串，避免中文显示问题
            table_data = [display_df.columns.tolist()]
            for _, row in display_df.iterrows():
                table_data.append([str(cell)[:20] + "..." if len(str(cell)) > 20 else str(cell) for cell in row.tolist()])
            
            # 创建表格
            table = Table(table_data, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0068B5')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(table)
            
            # 数据统计信息
            if len(df) > max_rows or len(df.columns) > max_cols:
                story.append(Paragraph(f"注: 完整数据包含 {len(df)} 行 × {len(df.columns)} 列，此处仅显示前 {max_rows} 行 × {max_cols} 列", style_dict["normal_style"]))
            
            story.append(Spacer(1, 10))
    
    def _add_alternatives_info(self, story: List, msg: Dict, style_dict: Dict):
        """添加其他可能的理解方式信息 - 修复视觉重叠问题"""
        # 其他可能的理解方式 (如果有)
        if msg.get("alternatives"):
            alternatives = msg["alternatives"]
            if alternatives and len(alternatives) > 0:
                # 添加足够的间距以避免与上一个内容重叠
                story.append(Spacer(1, 50))  # 50像素间距
                story.append(Paragraph(f"🤔 其他可能的理解方式 ({len(alternatives)}种)", style_dict["alternatives_heading_style"]))
                story.append(Spacer(1, 12))  # 标题后间距
                
                for j, alt in enumerate(alternatives[:3]):  # 限制显示前3个
                    alt_desc = alt.get("natural_description", alt.get("description", "无描述"))
                    confidence = alt.get("confidence", 0)
                    alt_text = f"{j+1}. {alt_desc} (置信度: {confidence:.2f})"
                    if hasattr(self, 'chinese_font_available') and self.chinese_font_available:
                        story.append(Paragraph(alt_text, style_dict["alternatives_style"]))
                    else:
                        story.append(Paragraph(f"{j+1}. Alternative interpretation (Confidence: {confidence:.2f})", style_dict["alternatives_style"]))
                    story.append(Spacer(1, 6))  # 每个选项后的间距
                
                if len(alternatives) > 3:
                    story.append(Paragraph(f"... 还有 {len(alternatives) - 3} 种其他理解方式", style_dict["alternatives_style"]))
                    story.append(Spacer(1, 10))
        
        # 选中的可能性 (如果有)
        if msg.get("selected_possibility"):
            selected = msg["selected_possibility"]
            # 添加间距以与其他理解方式分离
            story.append(Spacer(1, 30))
            story.append(Paragraph("✅ 选中的理解方式", style_dict["selected_heading_style"]))
            story.append(Spacer(1, 8))
            
            selected_desc = selected.get("natural_description", selected.get("description", "无描述"))
            confidence = selected.get("confidence", 0)
            selected_text = f"描述: {selected_desc}\n置信度: {confidence:.2f}"
            if hasattr(self, 'chinese_font_available') and self.chinese_font_available:
                story.append(Paragraph(selected_text, style_dict["selected_style"]))
            else:
                story.append(Paragraph(f"Selected interpretation (Confidence: {confidence:.2f})", style_dict["selected_style"]))
            story.append(Spacer(1, 15))
            selected_text = f"描述: {selected_desc}\n置信度: {confidence:.2f}"
            if hasattr(self, 'chinese_font_available') and self.chinese_font_available:
                story.append(Paragraph(selected_text, style_dict["normal_style"]))
            else:
                story.append(Paragraph(f"Selected interpretation (Confidence: {confidence:.2f})", style_dict["normal_style"]))
            story.append(Spacer(1, 10))
    
    def export_session_to_docx(self, session_data: Dict, session_title: str = "分析报告") -> Optional[str]:
        """导出完整会话为DOCX文档，包含所有对话内容、图表、AI思考过程等"""
        if not DOCX_AVAILABLE:
            return None
        
        try:
            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"DeepInsight_Complete_Report_{timestamp}.docx"
            filepath = os.path.join(self.exports_dir, filename)
            
            # 创建DOCX文档
            doc = Document()
            
            # 设置文档样式
            self._setup_docx_styles(doc)
            
            # 标题页
            title = doc.add_heading("Intel® DeepInsight 智能分析报告", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f"会话标题: {session_title}")
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
            doc.add_paragraph(f"系统版本: Intel® DeepInsight Pro v2.0")
            doc.add_paragraph("")
            
            # 目录概览
            doc.add_heading("📋 报告内容概览", 1)
            messages = session_data.get("messages", [])
            question_count = len([msg for msg in messages if msg["role"] == "user"])
            
            overview_items = [
                f"• 用户问题数量: {question_count}",
                f"• AI回答数量: {len([msg for msg in messages if msg['role'] == 'assistant'])}",
                f"• 包含SQL查询: {len([msg for msg in messages if msg.get('sql')])}",
                f"• 包含AI思考过程: {len([msg for msg in messages if msg.get('thought')])}"
            ]
            
            for item in overview_items:
                doc.add_paragraph(item)
            
            doc.add_paragraph("")
            
            # 处理每个对话
            question_num = 1
            for i, msg in enumerate(messages):
                if msg["role"] == "user":
                    # 用户问题
                    doc.add_heading(f"🙋‍♂️ 问题 {question_num}: {msg['content']}", 2)
                    question_num += 1
                
                elif msg["role"] == "assistant":
                    # AI思考过程 (如果有)
                    if msg.get("thought"):
                        doc.add_heading("🧠 AI思考过程 (DeepSeek R1 推理)", 3)
                        thought_content = msg["thought"][:1500] + "..." if len(msg["thought"]) > 1500 else msg["thought"]
                        thought_para = doc.add_paragraph(thought_content)
                        self._apply_thought_style(thought_para)
                    
                    # 表选择过程信息 (如果有)
                    self._add_table_selection_to_docx(doc, msg.get("table_selection_info"))
                    
                    # SQL查询 (如果有)
                    if msg.get("sql"):
                        doc.add_heading("💻 生成的SQL查询", 3)
                        sql_para = doc.add_paragraph(msg["sql"])
                        self._apply_code_style(sql_para)
                    
                    # 查询结果和数据
                    if "data" in msg and msg["data"]:
                        doc.add_heading("📊 查询结果", 3)
                        self._add_data_table_to_docx(doc, msg["data"])
                        
                        # 如果有图表数据，添加图表
                        if msg.get("charts"):
                            for chart_data in msg["charts"]:
                                self._add_chart_to_docx(doc, chart_data)
                    
                    # 商业洞察和分析
                    if msg.get("content"):
                        doc.add_heading("💡 商业洞察与分析", 3)
                        doc.add_paragraph(msg["content"])
                    
                    # 其他可能的理解方式和选中的可能性
                    self._add_alternatives_to_docx(doc, msg)
                    
                    # 添加分隔线
                    doc.add_paragraph("─" * 50)
            
            # 报告总结
            self._add_report_summary_to_docx(doc, question_count)
            
            # 保存文档
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            print(f"DOCX生成失败: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _setup_docx_styles(self, doc):
        """设置DOCX文档样式"""
        try:
            # 设置默认字体
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Microsoft YaHei'
            font.size = Pt(11)
        except:
            pass
    
    def _apply_thought_style(self, paragraph):
        """应用思考过程样式"""
        try:
            # 设置段落格式
            paragraph.style.font.name = 'Consolas'
            paragraph.style.font.size = Pt(9)
            # 添加边框效果（通过缩进模拟）
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.right_indent = Inches(0.5)
        except:
            pass
    
    def _apply_code_style(self, paragraph):
        """应用代码样式"""
        try:
            paragraph.style.font.name = 'Consolas'
            paragraph.style.font.size = Pt(9)
            paragraph.paragraph_format.left_indent = Inches(0.3)
        except:
            pass
    
    def _add_table_selection_to_docx(self, doc, table_info):
        """添加表选择过程信息到DOCX"""
        if not table_info or not any(table_info.values()):
            return
        
        doc.add_heading("🗄️ 智能表选择过程", 3)
        
        # 初步分析
        if table_info.get("initial_analysis"):
            doc.add_heading("第1步: 语义相似度初步筛选", 4)
            doc.add_paragraph(table_info["initial_analysis"])
        
        # Agent推理
        if table_info.get("agent_reasoning"):
            doc.add_heading("第2步: Agent智能筛选推理", 4)
            doc.add_paragraph(f"推理过程: {table_info['agent_reasoning']}")
        
        # 关联分析
        if table_info.get("join_analysis"):
            doc.add_heading("第3步: 表关联关系分析", 4)
            doc.add_paragraph(table_info["join_analysis"])
        
        # 最终选择结果
        if table_info.get("final_selection"):
            final_selection = table_info["final_selection"]
            selected_tables = final_selection.get("selected_tables", [])
            if selected_tables:
                doc.add_heading("最终选择的表:", 4)
                for table in selected_tables:
                    table_name = table.get("table_name", "未知表")
                    relevance_score = table.get("relevance_score", 0)
                    reasoning = table.get("reasoning", "无推理信息")
                    table_text = f"• {table_name} (相关度: {relevance_score:.2f}) - {reasoning}"
                    doc.add_paragraph(table_text)
    
    def _add_data_table_to_docx(self, doc, data):
        """添加数据表格到DOCX"""
        try:
            df = pd.DataFrame(data)
            if df.empty:
                return
            
            # 限制表格大小
            max_rows = 10
            max_cols = 6
            display_df = df.iloc[:max_rows, :max_cols]
            
            # 创建表格
            table = doc.add_table(rows=1, cols=len(display_df.columns))
            table.style = 'Table Grid'
            
            # 添加表头
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(display_df.columns):
                hdr_cells[i].text = str(column)
            
            # 添加数据行
            for _, row in display_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    cell_text = str(value)
                    if len(cell_text) > 30:
                        cell_text = cell_text[:27] + "..."
                    row_cells[i].text = cell_text
            
            # 数据统计信息
            if len(df) > max_rows or len(df.columns) > max_cols:
                doc.add_paragraph(f"注: 完整数据包含 {len(df)} 行 × {len(df.columns)} 列，此处仅显示前 {max_rows} 行 × {max_cols} 列")
        
        except Exception as e:
            print(f"添加数据表格到DOCX失败: {e}")
            doc.add_paragraph("[数据表格 - 处理失败]")
    
    def _add_chart_to_docx(self, doc, chart_data):
        """将图表添加到DOCX文档中"""
        try:
            img_path = self._convert_chart_to_image(chart_data)
            if img_path and os.path.exists(img_path):
                # 添加图表标题
                chart_title = chart_data.get("title", "数据图表")
                doc.add_heading(f"📊 {chart_title}", 4)
                
                # 使用绝对路径添加图片
                abs_img_path = os.path.abspath(img_path)
                doc.add_picture(abs_img_path, width=Inches(6))
                
                # 清理临时图片文件
                try:
                    os.remove(img_path)
                except:
                    pass
            else:
                # 如果图表转换失败，添加占位符
                doc.add_paragraph("📊 [图表数据 - 转换失败]")
        except Exception as e:
            print(f"添加图表到DOCX失败: {e}")
            doc.add_paragraph("📊 [图表数据 - 处理失败]")
    
    def _add_alternatives_to_docx(self, doc, msg):
        """添加其他可能的理解方式信息到DOCX"""
        # 其他可能的理解方式 (如果有)
        if msg.get("alternatives"):
            alternatives = msg["alternatives"]
            if alternatives and len(alternatives) > 0:
                doc.add_heading(f"🤔 其他可能的理解方式 ({len(alternatives)}种)", 3)
                for j, alt in enumerate(alternatives[:3]):  # 限制显示前3个
                    alt_desc = alt.get("natural_description", alt.get("description", "无描述"))
                    confidence = alt.get("confidence", 0)
                    alt_text = f"{j+1}. {alt_desc} (置信度: {confidence:.2f})"
                    doc.add_paragraph(alt_text)
                
                if len(alternatives) > 3:
                    doc.add_paragraph(f"... 还有 {len(alternatives) - 3} 种其他理解方式")
        
        # 选中的可能性 (如果有)
        if msg.get("selected_possibility"):
            selected = msg["selected_possibility"]
            doc.add_heading("✅ 选中的理解方式", 3)
            selected_desc = selected.get("natural_description", selected.get("description", "无描述"))
            confidence = selected.get("confidence", 0)
            selected_text = f"描述: {selected_desc}\n置信度: {confidence:.2f}"
            doc.add_paragraph(selected_text)
    
    def _add_report_summary_to_docx(self, doc, question_count):
        """添加报告总结到DOCX"""
        doc.add_heading("📈 报告总结", 1)
        doc.add_paragraph(f"本次会话共包含 {question_count} 个用户问题，系统通过Intel® DeepInsight智能分析引擎，结合OpenVINO™优化的语义理解和表选择算法，为每个问题提供了详细的分析过程和结果。")
        doc.add_paragraph("报告包含了完整的AI思考过程、表选择推理、SQL生成、数据分析结果以及商业洞察，为决策提供全面支持。")
        doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        doc.add_paragraph("技术支持: Intel® OpenVINO™ Toolkit")
    
    def _add_report_summary(self, story: List, question_count: int, style_dict: Dict):
        """添加报告总结"""
        story.append(Paragraph("📈 报告总结", style_dict["heading_style"]))
        story.append(Paragraph(f"本次会话共包含 {question_count} 个用户问题，系统通过Intel® DeepInsight智能分析引擎，结合OpenVINO™优化的语义理解和表选择算法，为每个问题提供了详细的分析过程和结果。", style_dict["normal_style"]))
        story.append(Paragraph("报告包含了完整的AI思考过程、表选择推理、SQL生成、数据分析结果以及商业洞察，为决策提供全面支持。", style_dict["normal_style"]))
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"报告生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}", style_dict["normal_style"]))
        story.append(Paragraph("技术支持: Intel® OpenVINO™ Toolkit", style_dict["normal_style"]))
    
    def export_data_to_excel(self, df: pd.DataFrame, filename_prefix: str = "data") -> str:
        """导出数据为Excel文件"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{filename_prefix}_{timestamp}.xlsx"
        filepath = os.path.join(self.exports_dir, filename)
        
        try:
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='数据', index=False)
                
                # 添加元数据工作表
                metadata = pd.DataFrame({
                    '属性': ['导出时间', '数据行数', '数据列数', '生成工具'],
                    '值': [
                        datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        len(df),
                        len(df.columns),
                        'Intel® DeepInsight'
                    ]
                })
                metadata.to_excel(writer, sheet_name='元数据', index=False)
            
            return filepath
        except Exception as e:
            print(f"Excel导出失败: {e}")
            return ""
    
    def export_data_to_csv(self, df: pd.DataFrame, filename_prefix: str = "data") -> str:
        """导出数据为CSV文件"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{filename_prefix}_{timestamp}.csv"
        filepath = os.path.join(self.exports_dir, filename)
        
        try:
            df.to_csv(filepath, index=False, encoding='utf-8-sig')
            return filepath
        except Exception as e:
            print(f"CSV导出失败: {e}")
            return ""
    
    def create_shareable_session(self, session_data: Dict, session_id: str) -> str:
        """创建可分享的会话快照"""
        try:
            share_id = str(uuid.uuid4())[:8]
            share_data = {
                "share_id": share_id,
                "original_session_id": session_id,
                "created_at": datetime.now().isoformat(),
                "session_data": session_data,
                "access_count": 0
            }
            
            share_file = os.path.join(self.shares_dir, f"share_{share_id}.json")
            with open(share_file, 'w', encoding='utf-8') as f:
                json.dump(share_data, f, indent=2, ensure_ascii=False)
            
            return share_id
        except Exception as e:
            print(f"创建分享失败: {e}")
            return ""
    
    def get_shared_session(self, share_id: str) -> Optional[Dict]:
        """获取分享的会话数据"""
        try:
            share_file = os.path.join(self.shares_dir, f"share_{share_id}.json")
            if not os.path.exists(share_file):
                return None
            
            with open(share_file, 'r', encoding='utf-8') as f:
                share_data = json.load(f)
            
            # 增加访问计数
            share_data["access_count"] += 1
            share_data["last_accessed"] = datetime.now().isoformat()
            
            with open(share_file, 'w', encoding='utf-8') as f:
                json.dump(share_data, f, indent=2, ensure_ascii=False)
            
            return share_data["session_data"]
        except Exception as e:
            print(f"获取分享会话失败: {e}")
            return None
    
    def get_download_link(self, filepath: str) -> str:
        """生成文件下载链接"""
        if not os.path.exists(filepath):
            return ""
        
        try:
            with open(filepath, "rb") as f:
                bytes_data = f.read()
            
            b64 = base64.b64encode(bytes_data).decode()
            filename = os.path.basename(filepath)
            
            if filepath.endswith('.pdf'):
                mime_type = "application/pdf"
            elif filepath.endswith('.xlsx'):
                mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            elif filepath.endswith('.csv'):
                mime_type = "text/csv"
            else:
                mime_type = "application/octet-stream"
            
            return f"data:{mime_type};base64,{b64}"
        except Exception as e:
            print(f"生成下载链接失败: {e}")
            return ""
    
    def cleanup_old_exports(self, days_old: int = 7):
        """清理旧的导出文件"""
        try:
            import time
            current_time = time.time()
            cutoff_time = current_time - (days_old * 24 * 60 * 60)
            
            for filename in os.listdir(self.exports_dir):
                filepath = os.path.join(self.exports_dir, filename)
                if os.path.isfile(filepath):
                    file_time = os.path.getmtime(filepath)
                    if file_time < cutoff_time:
                        os.remove(filepath)
                        print(f"已清理旧文件: {filename}")
        except Exception as e:
            print(f"清理文件失败: {e}")

# 全局实例

export_manager = ExportManager()
